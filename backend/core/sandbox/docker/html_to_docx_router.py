#!/usr/bin/env python3
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional
import re
from io import BytesIO
import copy
import base64
import tempfile
import os

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
except ImportError as e:
    raise ImportError(f"python-docx is not installed. Please install it with: pip install python-docx. Error: {e}")

from bs4 import BeautifulSoup


router = APIRouter(prefix="/document", tags=["docx-conversion"])

output_dir = Path("generated_docx")
output_dir.mkdir(exist_ok=True)


class ConvertRequest(BaseModel):
    doc_path: str = Field(..., description="Path to the document file (.doc for TipTap documents)")
    download: bool = Field(False, description="If true, returns the DOCX file directly. If false, returns JSON with download URL.")


class ConvertResponse(BaseModel):
    success: bool
    message: str
    docx_url: Optional[str] = None
    filename: Optional[str] = None


class HTMLToDocxConverter:
    def __init__(self, doc_path: str):
        self.doc_path = Path(doc_path)
        self.doc_data = {}
        self.document = None
        self.heading_style_map = {
            'h1': 'Heading 1',
            'h2': 'Heading 2',
            'h3': 'Heading 3',
            'h4': 'Heading 4',
            'h5': 'Heading 5',
            'h6': 'Heading 6'
        }

    def _setup_document_styles(self):
        """Configure professional document styles inspired by world-class document formatting."""
        section = self.document.sections[0]
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(1.2)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

        styles = self.document.styles

        # Enhanced Normal style with professional typography
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(12)
        normal_font.color.rgb = RGBColor(0, 0, 0)  # Pure black for readability
        normal_style.paragraph_format.space_after = Pt(10)
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.line_spacing = 1.8  # Improved readability

        # Professional body text style
        if 'IrisBodyText' not in styles:
            body_style = styles.add_style('IrisBodyText', WD_STYLE_TYPE.PARAGRAPH)
            body_style.base_style = normal_style
            body_style.font.name = 'Calibri'
            body_style.font.size = Pt(12)
            body_style.font.color.rgb = RGBColor(0, 0, 0)  # Pure black
            body_style.font.bold = False  # Ensure not bold
            body_style.paragraph_format.space_after = Pt(10)
            body_style.paragraph_format.space_before = Pt(0)
            body_style.paragraph_format.line_spacing = 1.8
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            body_style = styles['IrisBodyText']
            body_style.font.color.rgb = RGBColor(0, 0, 0)  # Ensure black
            body_style.font.bold = False

        # Enhanced blockquote with professional styling
        if 'IrisBlockQuote' not in styles:
            block_style = styles.add_style('IrisBlockQuote', WD_STYLE_TYPE.PARAGRAPH)
            block_style.base_style = body_style
            block_style.font.italic = True
            block_style.font.size = Pt(11)
            block_style.font.color.rgb = RGBColor(50, 50, 50)  # Dark gray for quotes
            block_style.paragraph_format.left_indent = Inches(0.5)
            block_style.paragraph_format.right_indent = Inches(0.3)
            block_style.paragraph_format.space_before = Pt(12)
            block_style.paragraph_format.space_after = Pt(12)
            block_style.paragraph_format.line_spacing = 1.6

        # Professional code block styling
        if 'IrisCodeBlock' not in styles:
            code_style = styles.add_style('IrisCodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_style.base_style = body_style
            code_style.font.name = 'Courier New'
            code_style.font.size = Pt(10)
            code_style.font.color.rgb = RGBColor(0, 0, 0)  # Black for code
            code_style.font.bold = False
            code_style.paragraph_format.left_indent = Inches(0.4)
            code_style.paragraph_format.right_indent = Inches(0.4)
            code_style.paragraph_format.space_before = Pt(10)
            code_style.paragraph_format.space_after = Pt(10)
            code_style.paragraph_format.line_spacing = 1.4

        # Professional table header styling
        if 'IrisTableHeader' not in styles:
            table_header = styles.add_style('IrisTableHeader', WD_STYLE_TYPE.PARAGRAPH)
            table_header.font.name = 'Calibri'
            table_header.font.size = Pt(12)
            table_header.font.bold = True
            table_header.font.color.rgb = RGBColor(0, 0, 0)  # Pure black
            table_header.paragraph_format.space_after = Pt(6)
            table_header.paragraph_format.space_before = Pt(6)
            table_header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Professional subtitle styling
        if 'IrisSubtitle' not in styles:
            subtitle = styles.add_style('IrisSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle.font.name = 'Calibri'
            subtitle.font.size = Pt(12)
            subtitle.font.color.rgb = RGBColor(73, 80, 87)
            subtitle.paragraph_format.space_after = Pt(6)

        # Professional list styles
        if 'IrisListBullet' not in styles:
            list_bullet = styles.add_style('IrisListBullet', WD_STYLE_TYPE.PARAGRAPH)
            list_bullet.base_style = body_style
            list_bullet.font.name = 'Calibri'
            list_bullet.font.size = Pt(12)
            list_bullet.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            list_bullet.font.bold = False
            list_bullet.paragraph_format.space_after = Pt(8)
            list_bullet.paragraph_format.space_before = Pt(0)
            list_bullet.paragraph_format.left_indent = Inches(0.3)
            list_bullet.paragraph_format.line_spacing = 1.6

        if 'IrisListNumber' not in styles:
            list_number = styles.add_style('IrisListNumber', WD_STYLE_TYPE.PARAGRAPH)
            list_number.base_style = body_style
            list_number.font.name = 'Calibri'
            list_number.font.size = Pt(12)
            list_number.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            list_number.font.bold = False
            list_number.paragraph_format.space_after = Pt(8)
            list_number.paragraph_format.space_before = Pt(0)
            list_number.paragraph_format.left_indent = Inches(0.3)
            list_number.paragraph_format.line_spacing = 1.6

        # Professional heading hierarchy with improved typography
        heading_styles = {
            'Heading 1': {
                'color': RGBColor(0, 0, 0),  # Pure black
                'size': Pt(26),
                'space_before': Pt(24),
                'space_after': Pt(14),
                'bold': True
            },
            'Heading 2': {
                'color': RGBColor(0, 0, 0),  # Pure black
                'size': Pt(22),
                'space_before': Pt(20),
                'space_after': Pt(12),
                'bold': True
            },
            'Heading 3': {
                'color': RGBColor(0, 0, 0),  # Pure black
                'size': Pt(18),
                'space_before': Pt(16),
                'space_after': Pt(10),
                'bold': True
            },
            'Heading 4': {
                'color': RGBColor(0, 0, 0),  # Pure black
                'size': Pt(16),
                'space_before': Pt(14),
                'space_after': Pt(8),
                'bold': True
            },
            'Heading 5': {
                'color': RGBColor(0, 0, 0),  # Pure black
                'size': Pt(14),
                'space_before': Pt(12),
                'space_after': Pt(6),
                'bold': True
            },
            'Heading 6': {
                'color': RGBColor(0, 0, 0),  # Pure black
                'size': Pt(13),
                'space_before': Pt(10),
                'space_after': Pt(6),
                'bold': True
            }
        }

        for heading_name, style_config in heading_styles.items():
            if heading_name in styles:
                heading_style = styles[heading_name]
                heading_style.font.name = 'Calibri'
                heading_style.font.color.rgb = style_config['color']
                heading_style.font.size = style_config['size']
                heading_style.font.bold = style_config['bold']
                heading_style.paragraph_format.space_before = style_config['space_before']
                heading_style.paragraph_format.space_after = style_config['space_after']
                heading_style.paragraph_format.line_spacing = 1.2

    def _add_document_header(self):
        title = self.doc_data.get('title', 'Document')
        
        # Professional title styling - use proper Heading 1 style
        title_paragraph = self.document.add_paragraph(title, style='Heading 1')
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(12)

        # Professional date styling
        date_paragraph = self.document.add_paragraph(
            datetime.now().strftime('%B %d, %Y'),
            style='IrisSubtitle'
        )
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.paragraph_format.space_after = Pt(6)

        # Professional branding
        created_by_paragraph = self.document.add_paragraph(
            'Created by Iris Intelligence Exclusively for You',
            style='IrisSubtitle'
        )
        created_by_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        created_by_paragraph.paragraph_format.space_after = Pt(18)

        # Professional separator line
        separator = self.document.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator_run = separator.add_run('─' * 50)
        separator_run.font.color.rgb = RGBColor(229, 231, 235)  # Very light gray
        separator_run.font.size = Pt(10)
        separator.paragraph_format.space_before = Pt(6)
        separator.paragraph_format.space_after = Pt(12)

    def _get_heading_color(self, level):
        """Get professional heading color based on level."""
        # All headings use pure black for maximum readability
        return RGBColor(0, 0, 0)

    def _get_heading_size(self, level):
        """Get professional heading size based on level."""
        sizes = {
            1: Pt(26),
            2: Pt(22),
            3: Pt(18),
            4: Pt(16),
            5: Pt(14),
            6: Pt(13)
        }
        return sizes.get(level, Pt(18))
        
    def load_document(self):
        if not self.doc_path.exists():
            raise FileNotFoundError(f"Document not found: {self.doc_path}")
            
        with open(self.doc_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        try:
            self.doc_data = json.loads(content)
            if self.doc_data.get('type') != 'tiptap_document':
                raise ValueError("Not a valid TipTap document")
        except json.JSONDecodeError:
            self.doc_data = {
                'title': self.doc_path.stem,
                'content': content,
                'metadata': {}
            }
    
    def create_docx(self) -> Document:
        self.document = Document()
        self._setup_document_styles()
        self._add_document_header()

        html_content = self.doc_data.get('content', '')
        self.convert_html_to_docx(html_content)
        
        return self.document
    
    def convert_html_to_docx(self, html_content: str):
        if not html_content:
            return
            
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for element in soup.contents:
            if getattr(element, 'name', None) is None and not str(element).strip():
                continue
            self.process_element(element)
    
    def apply_run_formatting(self, run, formatting: Dict):
        if formatting.get('bold') is not None:
            run.bold = formatting['bold']
        if formatting.get('italic') is not None:
            run.italic = formatting['italic']
        if formatting.get('underline') is not None:
            run.underline = formatting['underline']
        if formatting.get('strike') is not None:
            run.font.strike = formatting['strike']
        if formatting.get('font_name'):
            run.font.name = formatting['font_name']
        if formatting.get('font_size'):
            run.font.size = formatting['font_size']
        if formatting.get('color'):
            run.font.color.rgb = formatting['color']

    def _set_cell_background(self, cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def process_element(self, element, parent_paragraph=None):
        if getattr(element, 'name', None) is None:
            text = str(element)
            if not text or not text.strip():
                return
            if parent_paragraph is not None:
                run = parent_paragraph.add_run(text)
                self.apply_run_formatting(run, {})
            else:
                paragraph = self.document.add_paragraph(style='IrisBodyText')
                paragraph.paragraph_format.space_after = Pt(10)
                paragraph.paragraph_format.line_spacing = 1.8
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run(text.strip())
                run.font.bold = False
                run.font.color.rgb = RGBColor(0, 0, 0)
                self.apply_run_formatting(run, {})
            return

        tag = element.name.lower()

        if tag == 'p':
            paragraph = self.document.add_paragraph(style='IrisBodyText')
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.8
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for child in element.children:
                self.process_inline_element(child, paragraph, {'bold': False})

        elif tag in self.heading_style_map:
            style_name = self.heading_style_map[tag]
            paragraph = self.document.add_paragraph(style=style_name)
            
            # Process children with proper heading formatting
            for child in element.children:
                self.process_inline_element(child, paragraph, {})
            
            # If no children processed, add the text directly
            if not paragraph.text.strip():
                run = paragraph.add_run(element.get_text())
                # Apply heading formatting based on level
                heading_level = int(tag[1])  # Extract number from h1, h2, etc.
                if heading_level <= 3:
                    run.bold = True
                run.font.name = 'Calibri'
                run.font.color.rgb = self._get_heading_color(heading_level)
                run.font.size = self._get_heading_size(heading_level)

        elif tag == 'ul':
            for li in element.find_all('li', recursive=False):
                paragraph = self.document.add_paragraph(style='IrisListBullet')
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.left_indent = Inches(0.3)
                for child in li.children:
                    self.process_inline_element(child, paragraph, {})

        elif tag == 'ol':
            for li in element.find_all('li', recursive=False):
                paragraph = self.document.add_paragraph(style='IrisListNumber')
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.left_indent = Inches(0.3)
                for child in li.children:
                    self.process_inline_element(child, paragraph, {})

        elif tag == 'blockquote':
            paragraph = self.document.add_paragraph(style='IrisBlockQuote')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)
            for child in element.children:
                self.process_inline_element(child, paragraph, {'italic': True})
            if not paragraph.text.strip():
                run = paragraph.add_run(element.get_text())
                self.apply_run_formatting(run, {'italic': True, 'color': RGBColor(73, 80, 87)})

        elif tag == 'pre':
            paragraph = self.document.add_paragraph(style='IrisCodeBlock')
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)
            
            # Add background shading for code blocks
            pPr = paragraph._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'F5F5F5')  # Light grey background
            pPr.append(shd)
            
            code_text = element.get_text()
            run = paragraph.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            run.font.bold = False

        elif tag == 'hr':
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(16)
            paragraph.paragraph_format.space_after = Pt(16)
            run = paragraph.add_run('─' * 60)
            run.font.color.rgb = RGBColor(229, 231, 235)  # Professional light gray
            run.font.size = Pt(12)

        elif tag == 'br':
            if parent_paragraph:
                parent_paragraph.add_run().add_break()
            else:
                spacer = self.document.add_paragraph(style='IrisBodyText')
                spacer.add_run('').add_break()

        elif tag == 'table':
            self.process_table(element)

        elif tag == 'img':
            alt_text = element.get('alt', 'Image')
            src = element.get('src', '')
            
            # Try to embed the actual image
            image_embedded = False
            if src:
                try:
                    image_data = None
                    
                    # Handle data URIs (base64 encoded images)
                    if src.startswith('data:image/'):
                        # Extract base64 data
                        match = re.match(r'data:image/[^;]+;base64,(.+)', src)
                        if match:
                            image_data = base64.b64decode(match.group(1))
                    
                    # Handle file paths
                    elif not src.startswith(('http://', 'https://', 'blob:')):
                        # Try to read from file system
                        file_path = Path(src)
                        if not file_path.is_absolute():
                            # Try relative to workspace
                            file_path = Path('/workspace') / src.lstrip('/')
                        
                        if file_path.exists() and file_path.is_file():
                            with open(file_path, 'rb') as f:
                                image_data = f.read()
                    
                    # If we have image data, embed it
                    if image_data:
                        # Create temporary file for the image
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.tmp') as temp_file:
                            temp_file.write(image_data)
                            temp_file_path = temp_file.name
                        
                        try:
                            # Add image to document
                            paragraph = self.document.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph.paragraph_format.space_before = Pt(12)
                            paragraph.paragraph_format.space_after = Pt(12)
                            
                            # Add the image with appropriate sizing
                            run = paragraph.add_run()
                            picture = run.add_picture(temp_file_path, width=Inches(5.5))
                            image_embedded = True
                            
                            # Add caption if alt text exists
                            if alt_text and alt_text.lower() != 'image':
                                caption_paragraph = self.document.add_paragraph(style='IrisBodyText')
                                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                caption_paragraph.paragraph_format.space_before = Pt(4)
                                caption_paragraph.paragraph_format.space_after = Pt(12)
                                caption_run = caption_paragraph.add_run(alt_text)
                                caption_run.italic = True
                                caption_run.font.size = Pt(10)
                                caption_run.font.color.rgb = RGBColor(107, 114, 128)
                        finally:
                            # Clean up temporary file
                            try:
                                os.unlink(temp_file_path)
                            except:
                                pass
                except Exception as e:
                    print(f"Failed to embed image: {e}")
                    pass
            
            # If image embedding failed, show placeholder
            if not image_embedded:
                paragraph = self.document.add_paragraph(style='IrisBodyText')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(12)
                
                img_run = paragraph.add_run(f'[Image: {alt_text}]')
                img_run.italic = True
                img_run.font.color.rgb = RGBColor(107, 114, 128)
                img_run.font.size = Pt(11)
                
                if src and not src.startswith('data:'):
                    ref_run = paragraph.add_run(f'\n({src})')
                    ref_run.font.size = Pt(9)
                    ref_run.font.color.rgb = RGBColor(156, 163, 175)
                    ref_run.font.italic = True

        else:
            for child in element.children:
                self.process_element(child, parent_paragraph)

    def process_inline_element(self, element, paragraph, formatting: Dict):
        if getattr(element, 'name', None) is None:
            text = str(element)
            if text:
                run = paragraph.add_run(text)
                self.apply_run_formatting(run, formatting)
            return

        tag = element.name.lower()
        new_formatting = copy.deepcopy(formatting)

        if tag in ['strong', 'b']:
            new_formatting['bold'] = True
        elif tag in ['em', 'i']:
            new_formatting['italic'] = True
        elif tag == 'u':
            new_formatting['underline'] = True
        elif tag in ['s', 'del', 'strike']:
            new_formatting['strike'] = True
        elif tag == 'code':
            new_formatting['font_name'] = 'Courier New'
            new_formatting['font_size'] = Pt(10)
            new_formatting['color'] = RGBColor(200, 0, 0)  # Red for inline code
        elif tag == 'span':
            pass
        elif tag == 'br':
            paragraph.add_run().add_break()
            return

        if tag == 'a':
            new_formatting['underline'] = True
            new_formatting['color'] = RGBColor(37, 99, 235)  # Professional blue
            for child in element.children:
                self.process_inline_element(child, paragraph, new_formatting)
            href = element.get('href', '')
            if href:
                reference_run = paragraph.add_run(f' ({href})')
                reference_run.font.size = Pt(9)
                reference_run.font.color.rgb = RGBColor(107, 114, 128)  # Professional gray
            return

        for child in element.children:
            self.process_inline_element(child, paragraph, new_formatting)
    
    def process_table(self, table_element):
        rows = table_element.find_all('tr')
        if not rows:
            return

        max_cols = 0
        for row in rows:
            cells = row.find_all(['td', 'th'])
            max_cols = max(max_cols, len(cells))

        if max_cols == 0:
            return

        table = self.document.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'  # Professional grid style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = True

        # Set professional table formatting
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Set professional cell padding
                cell._tc.get_or_add_tcPr().append(
                    self._create_cell_margins(Inches(0.1), Inches(0.1), Inches(0.05), Inches(0.05))
                )

        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            is_header_row = any(cell.name == 'th' for cell in cells)
            
            for j in range(max_cols):
                table_cell = table.rows[i].cells[j]
                table_cell.text = ''
                paragraph = table_cell.paragraphs[0]
                paragraph.style = 'IrisBodyText'
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.space_before = Pt(6)

                if j < len(cells):
                    cell = cells[j]
                    is_header = cell.name == 'th'
                    cell_formatting = {'bold': is_header}
                    
                    for child in cell.children:
                        self.process_inline_element(child, paragraph, cell_formatting)
                    if not paragraph.text.strip():
                        run = paragraph.add_run(cell.get_text().strip())
                        self.apply_run_formatting(run, cell_formatting)

                    # Professional header styling
                    if is_header:
                        paragraph.style = 'IrisTableHeader'
                        # Professional header background color
                        self._set_cell_background(table_cell, 'F8F9FA')
                        # Add professional border
                        self._set_cell_border(table_cell, '2F3437', 'single')
                    else:
                        # Professional data cell styling
                        if i % 2 == 1:  # Alternating row colors
                            self._set_cell_background(table_cell, 'FAFBFC')
                        self._set_cell_border(table_cell, 'E5E7EB', 'single')
                else:
                    paragraph.add_run('')

    def _create_cell_margins(self, top, right, bottom, left):
        """Create professional cell margins for table cells."""
        margins = OxmlElement('w:tcMar')
        
        for margin_name, margin_value in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), str(int(margin_value * 1440)))  # Convert to twips
            margin.set(qn('w:type'), 'dxa')
            margins.append(margin)
        
        return margins

    def _set_cell_border(self, cell, color_hex, border_style):
        """Set professional cell borders."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Remove existing borders
        for border_name in ['top', 'left', 'bottom', 'right']:
            existing_border = tcPr.find(f'./w:{border_name}', namespaces=nsmap)
            if existing_border is not None:
                parent = existing_border.getparent()
                (parent or tcPr).remove(existing_border)
        
        # Add new borders
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), border_style)
            border.set(qn('w:sz'), '4')  # Border width
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color_hex)
            tcPr.append(border)
    
    async def convert_to_docx(self, store_locally: bool = True) -> tuple:
        self.load_document()
        
        doc = self.create_docx()
        
        doc_title = self.doc_data.get('title', 'document')
        safe_title = re.sub(r'[^\w\s-]', '', doc_title.lower())
        safe_title = re.sub(r'[-\s]+', '-', safe_title)[:50]
        
        if store_locally:
            docx_path = output_dir / f"{safe_title}.docx"
            doc.save(str(docx_path))
            return docx_path, safe_title
        else:
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read(), safe_title


@router.post("/convert-to-docx")
async def convert_document_to_docx(request: ConvertRequest):
    """
    Convert TipTap HTML document to DOCX.
    
    Takes a document file path and returns either:
    - DOCX file directly (if download=true) - uses document title as filename
    - JSON response with download URL (if download=false, default)
    """
    try:
        # Validate document path exists
        doc_path = Path(request.doc_path)
        if not doc_path.exists():
            raise HTTPException(status_code=404, detail=f"Document not found: {request.doc_path}")
        
        # Create converter
        converter = HTMLToDocxConverter(request.doc_path)
        
        # If download is requested, return file directly
        if request.download:
            docx_content, doc_name = await converter.convert_to_docx(store_locally=False)
            
            return Response(
                content=docx_content,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=\"{doc_name}.docx\"",
                    "X-Daytona-Skip-Preview-Warning": "true",
                }
            )
        
        # Otherwise, store locally and return JSON with download URL
        docx_path, doc_name = await converter.convert_to_docx(store_locally=True)
        
        docx_url = f"/downloads/{docx_path.name}"
        
        return ConvertResponse(
            success=True,
            message=f"DOCX generated successfully",
            docx_url=docx_url,
            filename=docx_path.name
        )
        
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"❌ Conversion error: {e}")
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")


@router.get("/health")
async def docx_health_check():
    """DOCX service health check endpoint."""
    return {"status": "healthy", "service": "docx-converter"}
